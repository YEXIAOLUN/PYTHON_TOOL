#-*- coding:utf-8 -*-
from docx import Document
from docx.shared import Inches
import http.client
import hashlib
import urllib
import random
import json
from time import sleep
appid = '20190708000315838' #你的appid
secretKey = 'H09WF2yyqhMqHEyZUIvz' #你的密钥
def BaiduFanyi(q):
    httpClient = None 
    myurl = '/api/trans/vip/translate'
    fromLang = 'en'
    toLang = 'zh'
    salt = random.randint(32768, 65536)

    sign = appid+q+str(salt)+secretKey
    m2 = hashlib.md5()
    m2.update(sign.encode("utf-8"))
    sign = m2.hexdigest()
    myurl = myurl+'?appid='+appid+'&q='+urllib.parse.quote(q)+'&from='+fromLang+'&to='+toLang+'&salt='+str(salt)+'&sign='+sign
     
    try:
        httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
        httpClient.request('GET', myurl)
 
        #response是HTTPResponse对象
        response = httpClient.getresponse()
        j=json.loads(response.read())
        #print(j)
        txt=j['trans_result'][0]['dst']
        return txt
    except Exception as e:
        print (e)
    finally:
        if httpClient:
            httpClient.close()
def main():
    words=str(input('请输入word的文件名：'))
    wordssave=str(input('请输入word保存的文件名：'))
    document = Document(words)
    documentsave = Document(wordssave)
    for paragraph in document.paragraphs:
        fanyitxt=BaiduFanyi(paragraph.text)
        sleep(1)
        documentsave.add_paragraph(fanyitxt)
        print(fanyitxt)
        documentsave.save(wordssave)
if __name__ == '__main__':
    main()
        
